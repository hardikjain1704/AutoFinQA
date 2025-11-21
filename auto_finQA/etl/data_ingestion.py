# etl/data_ingestion.py

import os
import sys
import json
import re
import pandas as pd
import pdfplumber
from pathlib import Path
from typing import List, Any, Dict

from docx import Document as DocxDocument
from dotenv import load_dotenv
from pymongo import MongoClient

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_mongodb import MongoDBAtlasVectorSearch

from utils.model_loader import ModelLoader
from logger import GLOBAL_LOGGER as log
from exception.custom_exception import AutoFinQAException
from utils.config_loader import load_config


SUPPORTED_EXTENSIONS = {
    '.pdf', '.txt', '.docx', '.md', 
    '.csv', '.json', '.xls', '.xlsx'
}

class DocumentLoader:
    
    @staticmethod
    def load_documents(file_paths: List[Path]) -> List[Document]:
        docs = []
        for file_path in file_paths:
            try:
                extension = file_path.suffix.lower()
                if extension not in SUPPORTED_EXTENSIONS:
                    log.warning(f"Skipping unsupported file type: {file_path}")
                    continue
                
                log.info(f"Loading document: {file_path.name}")
                
                if extension == '.pdf':
                    docs.extend(DocumentLoader._load_pdf(file_path))
                elif extension == '.docx':
                    docs.extend(DocumentLoader._load_docx(file_path))
                elif extension in ['.txt', '.md']:
                    docs.extend(DocumentLoader._load_text(file_path))
                elif extension == '.csv':
                    docs.extend(DocumentLoader._load_csv(file_path))
                elif extension in ['.xls', '.xlsx']:
                    docs.extend(DocumentLoader._load_excel(file_path))
                elif extension == '.json':
                    docs.extend(DocumentLoader._load_json(file_path))
                    
            except Exception as e:
                raise AutoFinQAException(f"Failed to load document: {file_path}", e)
        return docs
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Fixes spaced text like 'T E S L A' -> 'TESLA'."""
        if not text: return ""
        clean = re.sub(r'(?<!\S)([A-Z])\s(?=[A-Z](?:\s|$))', r'\1', text)
        return clean.replace('\n', ' ').strip()

    @staticmethod
    def _is_header_row(row: List[Any]) -> bool:
        """
        Adaptive Logic: Determines if a row is a header based on text density.
        Works for Date headers (Q1-2025) AND Category headers (Region, Model).
        """
        if not row: return False
        
        # Filter out None/Empty strings to check content
        cleaned_row = [str(x).strip() for x in row if x and str(x).strip()]
        if not cleaned_row: return False

        # Regex for pure numbers, currency, percentages, or parentheses (negative)
        number_pattern = re.compile(r'^[\d,\.\$%\(\)\s\-]+$')
        
        num_count = 0
        text_count = 0
        
        for cell in cleaned_row:
            if number_pattern.match(cell):
                num_count += 1
            else:
                text_count += 1
        
        # Heuristic: 
        # If text_count >= num_count, it's likely a header (e.g., "Region", "Model", "Q1-2025").
        # Pure data rows usually have 1 text label and 5+ numbers.
        if text_count >= num_count:
            return True
        
        return False

    @staticmethod
    def _row_to_markdown(row: List[Any]) -> str:
        """Converts a list of values to a Markdown table row pipe format."""
        clean_cells = []
        for cell in row:
            # Handle None and newlines
            txt = str(cell).replace('\n', ' ').strip() if cell is not None else ""
            clean_cells.append(txt)
        return "| " + " | ".join(clean_cells) + " |"

    @staticmethod
    def _load_pdf(file_path: Path) -> List[Document]:
        docs = []
        file_name = file_path.name
        log.info(f"Extracting content from {file_name}...")
        
        # CRITICAL: Increased snap_tolerance to fix "Prod-uction" splits
        table_settings = {
            "vertical_strategy": "lines", 
            "horizontal_strategy": "text",
            "snap_tolerance": 5, 
            "intersection_tolerance": 5
        }

        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    
                    # 1. Extract Text for Context
                    # We grab the first 300 chars to catch titles like "Operational Summary" or Years
                    raw_text = page.extract_text() or ""
                    page_text = DocumentLoader._clean_text(raw_text)
                    context_head = page_text[:300] if len(page_text) > 300 else page_text
                    
                    # 2. Extract Tables
                    try:
                        tables = page.extract_tables(table_settings)
                    except Exception:
                        tables = page.extract_tables()

                    if not tables:
                        # Fallback to pure text if no tables found
                        if len(page_text) > 50:
                             metadata = {
                                 "source": file_name, 
                                 "page_number": page_num, 
                                 "doc_type": "text"
                             }
                             docs.append(Document(page_content=page_text, metadata=metadata))
                        continue

                    for table_id, table in enumerate(tables, 1):
                        if not table or len(table) < 2: continue
                        
                        # --- SMART HEADER DETECTION ---
                        header_row = None
                        start_index = 0
                        
                        # Check first 3 rows for the best header candidate
                        for i in range(min(3, len(table))):
                            if DocumentLoader._is_header_row(table[i]):
                                header_row = table[i]
                                start_index = i + 1
                                break
                        
                        # Fallback generic headers
                        if not header_row:
                            header_row = [f"Col_{i}" for i in range(len(table[0]))]
                            start_index = 0
                        else:
                            # Normalize header row
                            header_row = [str(h).replace('\n', ' ').strip() if h else "" for h in header_row]

                        # --- MARKDOWN CONSTRUCTION ---
                        # We include the header in EVERY chunk so the vector search sees column names (Dates)
                        markdown_header = DocumentLoader._row_to_markdown(header_row)
                        markdown_sep = "| " + " | ".join(["---"] * len(header_row)) + " |"
                        
                        for row_num, row_data in enumerate(table[start_index:], start_index):
                            
                            # Clean row
                            clean_row = [str(cell).replace('\n', ' ').strip() if cell is not None else "" for cell in row_data]
                            
                            # Skip empty rows or page number artifacts
                            if not any(clean_row): continue
                            if len(clean_row) == 1 and clean_row[0].isdigit(): continue

                            markdown_row = DocumentLoader._row_to_markdown(clean_row)
                            
                            # FINAL CHUNK CONTENT
                            # Structure: Context -> Table Header -> Specific Row
                            content = (
                                f"Context: {context_head}...\n"
                                f"Table (Page {page_num}):\n"
                                f"{markdown_header}\n"
                                f"{markdown_sep}\n"
                                f"{markdown_row}"
                            )

                            # Metadata extraction for re-ranking
                            # We try to capture the first column as a "Row Label" for metadata
                            row_label = clean_row[0] if clean_row else "Unknown"

                            metadata = {
                                "source": file_name,
                                "page_number": page_num,
                                "table_id": f"p{page_num}_t{table_id}",
                                "row_number": row_num,
                                "doc_type": "tabular",
                                "row_label": row_label
                            }
                            
                            docs.append(Document(page_content=content, metadata=metadata))

            log.info(f"Extracted {len(docs)} chunks from {file_name}.")
        except Exception as e:
            log.warning(f"Failed to extract content from {file_name}", error=str(e))
            
        return docs

    @staticmethod
    def _load_docx(file_path: Path) -> List[Document]:
        doc = DocxDocument(file_path)
        content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return [Document(page_content=content, metadata={"source": file_path.name, "doc_type": "text"})]

    @staticmethod
    def _load_text(file_path: Path) -> List[Document]:
        content = file_path.read_text(encoding="utf-8")
        return [Document(page_content=content, metadata={"source": file_path.name, "doc_type": "text"})]

    @staticmethod
    def _load_csv(file_path: Path) -> List[Document]:
        df = pd.read_csv(file_path)
        docs = []
        headers = df.columns.tolist()
        markdown_header = "| " + " | ".join(headers) + " |"
        markdown_sep = "| " + " | ".join(["---"] * len(headers)) + " |"

        for index, row in df.iterrows():
            # Convert row to markdown values
            vals = [str(x).replace('\n', ' ') for x in row.values]
            markdown_row = "| " + " | ".join(vals) + " |"
            
            content_str = f"Table Data:\n{markdown_header}\n{markdown_sep}\n{markdown_row}"
            
            metadata = {
                "source": file_path.name,
                "doc_type": "tabular",
                "row_number": index + 1
            }
            docs.append(Document(page_content=content_str, metadata=metadata))
        log.info(f"Loaded {len(docs)} rows from CSV: {file_path.name}")
        return docs
        
    @staticmethod
    def _load_excel(file_path: Path) -> List[Document]:
        df = pd.read_excel(file_path)
        docs = []
        headers = df.columns.tolist()
        markdown_header = "| " + " | ".join([str(h) for h in headers]) + " |"
        markdown_sep = "| " + " | ".join(["---"] * len(headers)) + " |"

        for index, row in df.iterrows():
            vals = [str(x).replace('\n', ' ') if pd.notna(x) else "" for x in row.values]
            markdown_row = "| " + " | ".join(vals) + " |"
            
            content_str = f"Table Data:\n{markdown_header}\n{markdown_sep}\n{markdown_row}"
            
            metadata = {
                "source": file_path.name,
                "doc_type": "tabular",
                "row_number": index + 1,
                "original_data": json.dumps({k: str(v) for k,v in row.to_dict().items() if pd.notna(v)}, default=str)
            }
            docs.append(Document(page_content=content_str, metadata=metadata))
        log.info(f"Loaded {len(docs)} rows from Excel: {file_path.name}")
        return docs

    @staticmethod
    def _load_json(file_path: Path) -> List[Document]:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        docs = []
        if isinstance(data, list):
            for i, record in enumerate(data):
                content_str = json.dumps(record, indent=2)
                metadata = {
                    "source": file_path.name,
                    "doc_type": "tabular",
                    "row_number": i + 1
                }
                docs.append(Document(page_content=content_str, metadata=metadata))
        elif isinstance(data, dict):
            content_str = json.dumps(data, indent=2)
            docs.append(Document(page_content=content_str, metadata={"source": file_path.name, "doc_type": "text"}))
        
        log.info(f"Loaded content from JSON: {file_path.name}")
        return docs


class DataIngestion:
    def __init__(self):
        load_dotenv()
        self.config = load_config()
        self.model_loader = ModelLoader() 
        self.embeddings = self.model_loader.load_embeddings()
        self._load_db_config()
        log.info("DataIngestion pipeline initialized.")

    def _load_db_config(self):
        self.mongo_uri = os.getenv("MONGO_URI")
        if not self.mongo_uri:
            raise AutoFinQAException("Missing environment variable: MONGO_URI", sys)
        try:
            mongo_config = self.config['mongodb']
            self.db_name = mongo_config['db_name']
            self.collection_name = mongo_config['collection_name']
            self.index_name = mongo_config['index_name']
        except KeyError as e:
            raise AutoFinQAException(f"Missing required key in config.yaml under 'mongodb': {e}", e)

    def _split_documents(self, docs: List[Document]) -> List[Document]:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        text_docs = []
        tabular_docs = []
        for doc in docs:
            if doc.metadata.get("doc_type") == "tabular":
                tabular_docs.append(doc)
            else:
                text_docs.append(doc)
        
        if text_docs:
            chunks = text_splitter.split_documents(text_docs)
            log.info(f"Split {len(text_docs)} text documents into {len(chunks)} chunks.")
        else:
            chunks = []
        
        final_docs = chunks + tabular_docs
        log.info(f"Total documents to ingest: {len(final_docs)} ({len(chunks)} text chunks + {len(tabular_docs)} tabular docs)")
        return final_docs
    
    def ingest_single_document(self, doc_path_str: str):
        try:
            log.info(f"Starting ingestion for single document: {doc_path_str}")
            doc_path = Path(doc_path_str)
            raw_documents = DocumentLoader.load_documents([doc_path])
            
            if not raw_documents:
                log.warning(f"No content could be loaded from {doc_path_str}.")
                return

            chunks = self._split_documents(raw_documents)
            
            client = MongoClient(self.mongo_uri)
            collection = client[self.db_name][self.collection_name]

            log.info("Adding documents to MongoDB Atlas Vector Search...")
            MongoDBAtlasVectorSearch.from_documents(
                documents=chunks,
                embedding=self.embeddings,
                collection=collection,
                index_name=self.index_name
            )
            log.info(f"Successfully added {len(chunks)} chunks from '{doc_path.name}'.")

        except Exception as e:
            raise AutoFinQAException(f"Single document ingestion failed for {doc_path_str}", e)

    def ingest_from_directory(self, data_dir: str):
        try:
            data_path = Path(data_dir)
            if not data_path.is_dir():
                raise AutoFinQAException(f"Provided data path is not a directory: {data_dir}", sys)
            
            all_files = [p for p in data_path.rglob("*") if p.is_file()]
            raw_documents = DocumentLoader.load_documents(all_files)
            
            if not raw_documents:
                log.warning("No new supported documents found to ingest.")
                return

            chunks = self._split_documents(raw_documents)

            client = MongoClient(self.mongo_uri)
            collection = client[self.db_name][self.collection_name]
            
            log.info(f"Ingesting {len(chunks)} total document chunks into MongoDB...")
            MongoDBAtlasVectorSearch.from_documents(
                documents=chunks,
                embedding=self.embeddings,
                collection=collection,
                index_name=self.index_name
            )
            log.info(f"Successfully added {len(chunks)} document chunks to MongoDB.")

        except Exception as e:
            raise AutoFinQAException("Data ingestion pipeline failed.", e)

if __name__ == '__main__':
    DATA_DIRECTORY = "data" 
    if not os.path.exists(DATA_DIRECTORY):
        os.makedirs(DATA_DIRECTORY)
        sys.exit(0)
    try:
        ingestion_pipeline = DataIngestion()
        ingestion_pipeline.ingest_from_directory(DATA_DIRECTORY)
    except AutoFinQAException as e:
        log.error(e)
        sys.exit(1)